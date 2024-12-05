from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import ns, OxmlElement
from docx.oxml.ns import qn

from formatting_config import *
import logging

logger = logging.getLogger("my_app_logger")
doc = None


def format_margins(document):
    sections = document.sections
    for section in sections:
        section.top_margin = MARGIN_SETTINGS['top']
        section.bottom_margin = MARGIN_SETTINGS['bottom']
        section.left_margin = MARGIN_SETTINGS['left']
        section.right_margin = MARGIN_SETTINGS['right']
        section.page_width = PAGE_SETTINGS['width']
        section.page_height = PAGE_SETTINGS['height']


def format_spacing(item_name, paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = SPACING_SETTINGS[item_name]['before']
    paragraph_format.space_after = SPACING_SETTINGS[item_name]['after']
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = SPACING_SETTINGS[item_name]['line_spacing']


def format_font(item_name, run, styles=None):
    if item_name in FONT_SETTINGS['size']:
        run.font.size = FONT_SETTINGS['size'][item_name]
    if item_name in FONT_SETTINGS['style']:
        run.bold = FONT_SETTINGS['style'][item_name]['bold']
        run.italic = FONT_SETTINGS['style'][item_name]['italic']
    if item_name in FONT_SETTINGS['color']:
        run.font.color.rgb = FONT_SETTINGS['color'][item_name]
    if styles:
        for style in styles:
            if style == 'bold':
                run.bold = True
            if style == 'italic':
                run.italic = True


def format_alignment(item_name, content):
    content.alignment = ALIGNMENT_SETTINGS[item_name]


def format_indentation(item_name, paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = INDENTATION_SETTINGS[item_name]['left']
    paragraph_format.right_indent = INDENTATION_SETTINGS[item_name]['right']
    paragraph_format.first_line_indent = INDENTATION_SETTINGS[item_name]['first_line']


def format_content(item_name, paragraph, run, styles=None):
    format_font(item_name, run, styles)
    format_alignment(item_name, paragraph)
    format_spacing(item_name, paragraph)


def add_name(name):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(name)
    format_content('name', paragraph, run)


def add_contact_info(contact_info):
    contact_info_str = FONT_SETTINGS['style']['contact_info']['delimiter'].join(contact_info.values())
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(contact_info_str)
    format_content('contact_info', paragraph, run)


def add_personal_info(personal_info):
    add_name(personal_info['name'])
    contact_info = {k: v for k, v in personal_info.items() if k != 'name'}
    add_contact_info(contact_info)


def add_section_heading(heading):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(heading)
    format_content('section_heading', paragraph, run)

    # Add bottom border to heading
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 4 half-points = 2 points
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_infocard_heading(title_contents, timeline):
    # Add a table with 1 row and 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Calculate available width for the table
    total_width = PAGE_SETTINGS['width'] - 2 * MARGIN_SETTINGS['left']
    first_col_width = TABLE_SETTINGS['event_table']['first_col_width']
    second_col_width = total_width - first_col_width + TABLE_SETTINGS['event_table']['second_col_extra_width']

    # Set column widths
    table.columns[0].width = first_col_width
    table.columns[1].width = second_col_width

    # Format first column
    cell_1 = table.cell(0, 0)
    cell_1.width = first_col_width
    cell_1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    paragraph_1 = cell_1.paragraphs[0]

    if title_contents['type'] == 'project':
        # title = f'{title_contents["domain"]}, {title_contents["name"]}, ({title_contents["additional_info"]})'
        title = f'{title_contents["domain"]}, {title_contents["name"]}'  # Without additional info
    elif title_contents['type'] == 'experience':
        # title = f'{title_contents["company"]}, {title_contents["role"]}, {title_contents["location"]}, ({title_contents["additional_info"]})'
        title = f'{title_contents["company"]}, {title_contents["role"]}, {title_contents["location"]}'  # Without additional info
    elif title_contents['type'] == 'education':
        title = f'{title_contents["university"]}, {title_contents["location"]}'

    run = paragraph_1.add_run(title)
    format_content('infocard_heading_title', paragraph_1, run)

    # Format second column
    cell_2 = table.cell(0, 1)
    cell_2.width = second_col_width
    cell_2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    paragraph_2 = cell_2.paragraphs[0]

    if timeline:
        run_2 = paragraph_2.add_run(timeline)
        format_content('infocard_heading_timeline', paragraph_2, run_2)


def add_infocard_subheading(subheading):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(subheading)
    format_content('infocard_subheading', paragraph, run)


def add_bullet_points(bullet_points):
    for bullet_point in bullet_points:
        paragraph = doc.add_paragraph()
        pPr = paragraph._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), "0")
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), "1")
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.insert(0, numPr)

        run = paragraph.add_run(bullet_point)
        format_content('bullet_points', paragraph, run)
        format_indentation('bullet_points', paragraph)


def add_information_card(item):
    if item['title_contents']:
        add_infocard_heading(item['title_contents'], item['timeline'])
    if item['subtitle_contents']:
        add_infocard_subheading(item['subtitle_contents'])
    if item['bullet_points']:
        add_bullet_points(item['bullet_points'])


def add_skills_table(skills):
    # Add a new table with 2 columns
    skills_table = doc.add_table(rows=0, cols=2)
    skills_table.autofit = False
    format_alignment('skills_table', skills_table)  # Center the table

    # Set column widths to span the full width of the page minus margins
    page_width = PAGE_SETTINGS['width'] - MARGIN_SETTINGS['left'] * 2  # page width minus margins
    col1_width = TABLE_SETTINGS['skills_table']['first_col_width']
    col2_width = page_width - col1_width

    # Add rows and set cell text and formatting
    for skill in skills:
        category = skill['category']
        tools = skill['tools']
        row = skills_table.add_row().cells
        category_cell, skills_cell = row

        # Set column widths
        category_cell.width = col1_width
        skills_cell.width = col2_width

        # Add category name in italic
        category_run = category_cell.paragraphs[0].add_run(category)
        category_run.italic = True
        format_font('skills', category_run)
        format_alignment('skills', category_cell.paragraphs[0])

        # Add skills
        skills_run = skills_cell.paragraphs[0].add_run(', '.join(tools))
        format_font('skills', skills_run)
        format_alignment('skills', category_cell.paragraphs[0])

        # Set row height
        tr = row[0]._element.getparent()
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(ns.qn('w:val'),
                     str(TABLE_SETTINGS['skills_table']['row_height'] * INCH_TO_TWIPS))  # 0.15 inches = 300 twips
        trHeight.set(ns.qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

        for cell in row:
            format_spacing('skills', cell.paragraphs[0])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Make table gridlines visible
    tbl = skills_table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '4')  # 4 half-points = 2 points
        border.set(ns.qn('w:space'), '0')
        border.set(ns.qn('w:color'), 'auto')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)


def add_skills_passage(items):
    paragraph = doc.add_paragraph()
    for item in items:
        category = item['category']
        tools = item['tools']
        category_run = paragraph.add_run(f'{category}: ')
        format_content('skills_category', paragraph, category_run)
        tools_run = paragraph.add_run(', '.join(tools) + '  ')
        format_content('skills_tools', paragraph, tools_run)


def add_section(section_name, items, type='infocard'):
    add_section_heading(section_name)
    if type == 'infocard':
        for item in items:
            add_information_card(item)
    elif type == 'table':
        add_skills_table(items)
    elif type == 'passage':
        add_skills_passage(items)


def create_document(resume_contents):
    # Create a new Document
    global doc
    doc = Document()

    # Set document font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_SETTINGS['font']

    format_margins(doc)

    add_personal_info(resume_contents['personal_info'])
    logger.info('added personal info')
    add_section('EDUCATION', resume_contents['educations'])
    logger.info('added educations')
    # add_section('SKILLS', resume_contents['skills'], type='table')
    # add_section('SKILLS', resume_contents['skills'], type='passage')
    logger.info('added skills')
    add_section('WORK EXPERIENCE', resume_contents['experiences'])
    logger.info('added experience')
    add_section('PROJECTS', resume_contents['projects'])
    logger.info('added proj')

    return doc
